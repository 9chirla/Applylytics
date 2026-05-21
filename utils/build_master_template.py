"""
Build MasterTemplate.docx with consistent UK CV section styling.

Every section uses the same pattern: optional {%p if %}, bordered heading paragraph,
content paragraph(s), {%p endif %}. Styles are copied once from the Naveen reference CV.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

# UK CV typography (from Naveen reference DOCX)
NAME_PT = 18
NAME_COLOR = RGBColor(0x1F, 0x38, 0x64)
CONTACT_PT = 9
BODY_PT = 9.5
JOB_TITLE_PT = 10


def _delete_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def _clear_document_body(doc: Document) -> None:
    for para in reversed(doc.paragraphs):
        _delete_paragraph(para)


def _copy_ppr(source: Paragraph, target: Paragraph) -> None:
    sp = source._element.pPr
    if sp is not None:
        target._element.insert(0, deepcopy(sp))


def _copy_rpr(source_run, target_run) -> None:
    sr = source_run._element.rPr
    if sr is not None:
        target_run._element.insert(0, deepcopy(sr))


class _StyleBook:
    """Paragraph prototypes copied from the reference CV."""

    def __init__(self, source_path: Path) -> None:
        ref = Document(str(source_path))
        self.name_p = ref.paragraphs[0]
        self.contact_p = ref.paragraphs[1]
        self.heading_p = ref.paragraphs[2]  # PROFILE
        self.body_p = ref.paragraphs[3]
        # First list-style bullet in reference
        self.list_p = next(
            (p for p in ref.paragraphs if p.style and p.style.name == "List Paragraph"),
            self.body_p,
        )
        self.job_title_p = ref.paragraphs[5]


class CvDocxTemplateBuilder:
    """Programmatic docxtpl shell — one code path for all section headings."""

    def __init__(self, style_source: Path) -> None:
        self.style_source = style_source
        self.styles = _StyleBook(style_source)

    def _para(self, doc: Document, text: str = "") -> Paragraph:
        p = doc.add_paragraph()
        if text:
            p.add_run(text)
        return p

    def _name_block(self, doc: Document) -> None:
        p = self._para(doc)
        run = p.add_run("{{ name }}")
        _copy_ppr(self.styles.name_p, p)
        _copy_rpr(self.styles.name_p.runs[0], run)
        run.bold = True
        run.font.size = Pt(NAME_PT)
        if run.font.color is None or run.font.color.rgb is None:
            run.font.color.rgb = NAME_COLOR

    def _contact_block(self, doc: Document) -> None:
        p = self._para(doc)
        run = p.add_run(
            "{{ contact_phone }} • {{ contact_email }} • {{ contact_location }} • {{ linkedin }}"
        )
        _copy_ppr(self.styles.contact_p, p)
        _copy_rpr(self.styles.contact_p.runs[0], run)
        run.font.size = Pt(CONTACT_PT)

    def _section_heading(self, doc: Document, title: str) -> None:
        """Bold heading + bottom rule (same as PROFILE / SKILLS in reference)."""
        p = self._para(doc)
        run = p.add_run(title)
        _copy_ppr(self.styles.heading_p, p)
        _copy_rpr(self.styles.heading_p.runs[0], run)
        run.bold = True

    def _body(self, doc: Document, text: str) -> None:
        p = self._para(doc, text)
        if self.styles.body_p.runs:
            for run in p.runs:
                _copy_rpr(self.styles.body_p.runs[0], run)
                run.font.size = Pt(BODY_PT)

    def _job_title_line(self, doc: Document, text: str) -> None:
        p = self._para(doc, text)
        if self.styles.job_title_p.runs:
            for run in p.runs:
                _copy_rpr(self.styles.job_title_p.runs[0], run)

    def _list_line(self, doc: Document, text: str) -> None:
        p = self._para(doc, text)
        if self.styles.list_p.style:
            p.style = self.styles.list_p.style
        if self.styles.list_p.runs:
            for run in p.runs:
                _copy_rpr(self.styles.list_p.runs[0], run)
                if self.styles.list_p.runs[0].font.size:
                    run.font.size = self.styles.list_p.runs[0].font.size
                else:
                    run.font.size = Pt(BODY_PT)

    def _p_if_open(self, doc: Document, condition: str) -> None:
        self._para(doc, f"{{%p if {condition} %}}")

    def _p_if_close(self, doc: Document) -> None:
        self._para(doc, "{%p endif %}")

    def build(self, dest: Path) -> None:
        doc = Document(str(self.style_source))
        _clear_document_body(doc)

        self._name_block(doc)
        self._contact_block(doc)

        self._section_heading(doc, "PROFILE")
        self._body(doc, "{{ personal_statement }}")

        self._p_if_open(doc, "show_work_experience")
        self._section_heading(doc, "WORK EXPERIENCE")
        self._para(doc, "{%p for exp in work_experience %}")
        self._job_title_line(
            doc,
            "{{ exp.job_title }} | {{ exp.employer }} | {{ exp.dates }}",
        )
        self._para(doc, "{%p for bullet in exp.bullets %}")
        self._list_line(doc, "{{ bullet }}")
        self._para(doc, "{%p endfor %}")
        self._para(doc, "{%p endfor %}")
        self._p_if_close(doc)

        self._p_if_open(doc, "show_projects")
        self._section_heading(doc, "PROJECTS")
        self._para(doc, "{%p for proj in projects %}")
        self._job_title_line(doc, "{{ proj.header_line }}")
        self._para(doc, "{% if proj.description %}")
        self._body(doc, "{{ proj.description }}")
        self._para(doc, "{% endif %}")
        self._para(doc, "{%p for tech in proj.technologies %}")
        self._list_line(doc, "{{ tech }}")
        self._para(doc, "{%p endfor %}")
        self._para(doc, "{%p endfor %}")
        self._p_if_close(doc)

        self._p_if_open(doc, "show_skills")
        self._section_heading(doc, "SKILLS")
        self._body(doc, "{{ skills_line }}")
        self._p_if_close(doc)

        self._p_if_open(doc, "show_education")
        self._section_heading(doc, "EDUCATION")
        self._body(doc, "{{ education_block }}")
        self._p_if_close(doc)

        self._p_if_open(doc, "show_certifications")
        self._section_heading(doc, "CERTIFICATIONS")
        self._body(doc, "{{ certifications_block }}")
        self._p_if_close(doc)

        dest.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(dest))


def build_master_template(source: Path | str, dest: Path | str) -> None:
    """Build MasterTemplate.docx (``source`` supplies styles only)."""
    CvDocxTemplateBuilder(Path(source)).build(Path(dest))


def create_master_template(dest: Path | str, style_source: Path | str | None = None) -> None:
    """Public entry: write template to ``dest``."""
    root = Path(__file__).resolve().parent.parent
    ref = Path(style_source) if style_source else root / "Naveen_Reddy_Chirla_CV_Buitelaar.docx"
    build_master_template(ref, dest)


if __name__ == "__main__":
    _root = Path(__file__).resolve().parent.parent
    create_master_template(_root / "templates" / "MasterTemplate.docx")
