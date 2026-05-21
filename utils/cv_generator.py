"""Generate DOCX from template and context."""
from __future__ import annotations

from collections.abc import Mapping
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docxtpl import DocxTemplate

from applylytics.cv.sections import compute_section_visibility, sanitize_projects


def write_distinct_cv_templates(out_dir: Path | str) -> None:
    """(Re)build four .docx shells: same Jinja keys, different layout and typography. No sample CV text."""
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    _build_chronological(out / "Chronological.docx")
    _build_modern(out / "Modern.docx")
    _build_creative(out / "Creative.docx")
    _build_two_column(out / "Two-Column.docx")


def _jinja_para(doc: Document, text: str, *, bold: bool = False, size_pt: float | None = None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size_pt is not None:
        r.font.size = Pt(int(size_pt))


def _section_heading(doc: Document, title: str, *, level: int = 2) -> None:
    """Word built-in heading styles render bold in typical themes."""
    doc.add_heading(title, level=level)


def _build_chronological(path: Path) -> None:
    d = Document()
    p = d.add_paragraph()
    r = p.add_run("{{ name }}")
    r.bold = True
    r.font.size = Pt(20)
    _jinja_para(d, "{{ phone }} | {{ email }} | {{ location }}")
    _jinja_para(d, "LinkedIn: {{ linkedin }}")
    d.add_paragraph()
    _section_heading(d, "PROFILE", level=2)
    _jinja_para(d, "{{ personal_statement }}")
    _section_heading(d, "WORK EXPERIENCE", level=2)
    _jinja_para(d, "{{ work_block }}")
    _section_heading(d, "EDUCATION", level=2)
    _jinja_para(d, "{{ education_block }}")
    _section_heading(d, "SKILLS", level=2)
    _jinja_para(d, "{{ skills_line }}")
    _section_heading(d, "CERTIFICATIONS", level=2)
    _jinja_para(d, "{{ certifications_block }}")
    _section_heading(d, "INTERESTS", level=2)
    _jinja_para(d, "{{ interests }}")
    _section_heading(d, "REFERENCES", level=2)
    _jinja_para(d, "{{ references }}")
    d.save(path)


def _build_modern(path: Path) -> None:
    d = Document()
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("{{ name }}")
    r.bold = True
    r.font.size = Pt(26)
    for line, size in (
        ("{{ phone }}  ·  {{ email }}", 11),
        ("{{ location }}", 11),
        ("{{ linkedin }}", 10),
    ):
        p2 = d.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p2.add_run(line)
        rr.font.size = Pt(size)
    d.add_paragraph()
    _section_heading(d, "Professional profile", level=1)
    _jinja_para(d, "{{ personal_statement }}")
    _section_heading(d, "Experience", level=1)
    _jinja_para(d, "{{ work_block }}")
    _section_heading(d, "Education", level=1)
    _jinja_para(d, "{{ education_block }}")
    _section_heading(d, "Skills", level=1)
    _jinja_para(d, "{{ skills_line }}")
    _section_heading(d, "Certifications", level=1)
    _jinja_para(d, "{{ certifications_block }}")
    _section_heading(d, "Interests", level=1)
    _jinja_para(d, "{{ interests }}")
    _section_heading(d, "References", level=1)
    _jinja_para(d, "{{ references }}")
    d.save(path)


def _build_creative(path: Path) -> None:
    d = Document()
    accent = RGBColor(0x5C, 0x6E, 0xF8)
    p = d.add_paragraph()
    r = p.add_run("{{ name }}")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = accent
    _jinja_para(d, "{{ phone }}  ·  {{ email }}  ·  {{ location }}")
    _jinja_para(d, "{{ linkedin }}", size_pt=10)
    d.add_paragraph()

    def band(title: str, body: str) -> None:
        h = d.add_paragraph()
        hr = h.add_run(title.upper())
        hr.bold = True
        hr.font.size = Pt(13)
        hr.font.color.rgb = accent
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        d.add_paragraph(body)

    band("Profile", "{{ personal_statement }}")
    band("Experience", "{{ work_block }}")
    band("Education", "{{ education_block }}")
    band("Skills", "{{ skills_line }}")
    band("Certifications", "{{ certifications_block }}")
    band("Interests", "{{ interests }}")
    band("References", "{{ references }}")
    d.save(path)


def _build_two_column(path: Path) -> None:
    d = Document()
    table = d.add_table(rows=1, cols=2)
    table.autofit = False
    try:
        table.columns[0].width = Inches(2.35)
        table.columns[1].width = Inches(4.65)
    except Exception:
        pass
    left, right = table.rows[0].cells[0], table.rows[0].cells[1]

    def cell_heading(cell, title: str) -> None:
        p = cell.add_paragraph()
        rr = p.add_run(title)
        rr.bold = True
        rr.font.size = Pt(11)

    p0 = left.paragraphs[0]
    rr0 = p0.add_run("{{ name }}")
    rr0.bold = True
    rr0.font.size = Pt(16)
    left.add_paragraph("{{ phone }}")
    left.add_paragraph("{{ email }}")
    left.add_paragraph("{{ location }}")
    left.add_paragraph("{{ linkedin }}")
    left.add_paragraph()
    cell_heading(left, "SKILLS")
    left.add_paragraph("{{ skills_line }}")
    left.add_paragraph()
    cell_heading(left, "EDUCATION")
    left.add_paragraph("{{ education_block }}")
    left.add_paragraph()
    cell_heading(left, "CERTIFICATIONS")
    left.add_paragraph("{{ certifications_block }}")

    rp0 = right.paragraphs[0]
    rrh = rp0.add_run("PROFILE")
    rrh.bold = True
    rrh.font.size = Pt(12)
    right.add_paragraph("{{ personal_statement }}")
    right.add_paragraph()
    h2 = right.add_paragraph()
    r2 = h2.add_run("EXPERIENCE")
    r2.bold = True
    r2.font.size = Pt(12)
    right.add_paragraph("{{ work_block }}")
    right.add_paragraph()
    h3 = right.add_paragraph()
    r3 = h3.add_run("INTERESTS & REFERENCES")
    r3.bold = True
    r3.font.size = Pt(12)
    right.add_paragraph("{{ interests }}")
    right.add_paragraph()
    right.add_paragraph("{{ references }}")
    d.save(path)


def _projects_for_template(data: Mapping) -> list[dict[str, Any]]:
    """Structured projects with header_line for DOCX loops (placeholders stripped)."""
    projects = sanitize_projects(data.get("projects"))
    for proj in projects:
        name = proj["name"]
        date = proj.get("date") or ""
        proj["header_line"] = f"{name} | {date}" if date else name
    return projects


def cv_context_from_structured(
    data: Mapping,
    source_resume_text: str | None = None,
) -> dict[str, Any]:
    """Flatten optimised CV JSON into template fields; sections driven by real content + source CV."""
    w_lines: list[str] = []
    work_experience: list[dict[str, str | list[str]]] = []
    for job in data.get("work_experience") or []:
        if not isinstance(job, dict):
            continue
        bullets = [str(b).strip() for b in (job.get("bullets") or []) if str(b).strip()]
        work_experience.append(
            {
                "job_title": str(job.get("job_title", "")).strip(),
                "employer": str(job.get("employer", "")).strip(),
                "dates": str(job.get("dates", "")).strip(),
                "bullets": bullets,
            }
        )
        w_lines.append(
            f"{str(job.get('job_title', '')).strip()} | {str(job.get('employer', '')).strip()} | {str(job.get('dates', '')).strip()}"
        )
        for b in job.get("bullets") or []:
            w_lines.append(f"- {b}")
        w_lines.append("")
    e_lines: list[str] = []
    for edu in data.get("education") or []:
        if not isinstance(edu, dict):
            continue
        g = edu.get("grade")
        gpart = f" | {g}" if g else ""
        e_lines.append(
            f"{str(edu.get('degree', '')).strip()} | {str(edu.get('institution', '')).strip()} | {str(edu.get('dates', '')).strip()}{gpart}"
        )
    certs = [
        str(c).lstrip("-• ").strip()
        for c in (data.get("certifications") or [])
        if str(c).lstrip("-• ").strip()
    ]
    cert_blk = "\n".join(f"- {c}" for c in certs)
    projects = _projects_for_template(data)
    visibility = compute_section_visibility(data, source_resume_text)

    skills_only = data.get("skills") or []
    skills_line = " • ".join(str(s).strip() for s in skills_only if str(s).strip())
    interests_raw = data.get("interests")
    if isinstance(interests_raw, list):
        interests = ", ".join(str(i).strip() for i in interests_raw if str(i).strip())
    else:
        interests = str(interests_raw or "")
    phone = str(data.get("contact_phone", ""))
    email = str(data.get("contact_email", ""))
    location = str(data.get("contact_location", ""))
    return {
        "name": str(data.get("name", "")),
        "phone": phone,
        "email": email,
        "location": location,
        "contact_phone": phone,
        "contact_email": email,
        "contact_location": location,
        "linkedin": str(data.get("linkedin", "")),
        "personal_statement": str(data.get("personal_statement", "")),
        "work_experience": work_experience,
        "work_block": "\n".join(w_lines).strip(),
        "show_work_experience": visibility["show_work_experience"],
        "projects": projects,
        "show_projects": visibility["show_projects"],
        "education_block": "\n".join(e_lines).strip(),
        "show_education": visibility["show_education"],
        "skills_line": skills_line,
        "show_skills": visibility["show_skills"],
        "certifications_block": cert_blk,
        "show_certifications": visibility["show_certifications"],
        "interests": interests,
        "show_interests": visibility["show_interests"],
        "references": str(data.get("references", "") or "Available on request"),
    }


def render_cv(template_path: str, output_path: str, context: dict) -> str:
    """Fill a DOCX template with context dict and save to output_path."""
    from applylytics.cv.docx_normalize import normalize_rendered_docx

    template = DocxTemplate(template_path)
    template.render(context)
    template.save(output_path)
    normalize_rendered_docx(output_path)
    return output_path


if __name__ == "__main__":
    write_distinct_cv_templates(Path(__file__).resolve().parent.parent / "templates")
