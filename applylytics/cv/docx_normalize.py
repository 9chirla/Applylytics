"""Post-render DOCX normalisation — safety net for section heading styles."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

SECTION_TITLES = frozenset(
    {
        "PROFILE",
        "WORK EXPERIENCE",
        "PROJECTS",
        "SKILLS",
        "EDUCATION",
        "CERTIFICATIONS",
        "INTERESTS",
        "REFERENCES",
    }
)


def _load_heading_prototype() -> Paragraph | None:
    root = Path(__file__).resolve().parent.parent.parent
    ref_path = root / "Naveen_Reddy_Chirla_CV_Buitelaar.docx"
    if not ref_path.is_file():
        return None
    ref = Document(str(ref_path))
    for p in ref.paragraphs:
        if p.text.strip() == "PROFILE":
            return p
    return None


def _apply_heading_style(para: Paragraph, proto: Paragraph) -> None:
    title = para.text.strip().split("\n", 1)[0].strip().upper()
    if title not in SECTION_TITLES:
        return
    para.clear()
    run = para.add_run(title)
    sp = proto._element.pPr
    if sp is not None:
        para._element.insert(0, deepcopy(sp))
    if proto.runs:
        sr = proto.runs[0]._element.rPr
        if sr is not None:
            run._element.insert(0, deepcopy(sr))
    run.bold = True


def _remove_empty_paragraphs(doc: Document) -> None:
    for para in reversed(doc.paragraphs):
        if not (para.text or "").strip():
            para._element.getparent().remove(para._element)


def normalize_rendered_docx(doc_path: str) -> None:
    """After docxtpl render: blank paragraph removal + bordered section headings."""
    doc = Document(doc_path)
    proto = _load_heading_prototype()
    if proto is not None:
        for para in doc.paragraphs:
            first = (para.text or "").strip().split("\n", 1)[0].strip().upper()
            if first in SECTION_TITLES:
                _apply_heading_style(para, proto)
    _remove_empty_paragraphs(doc)
    doc.save(doc_path)
