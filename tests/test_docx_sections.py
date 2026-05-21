"""Section heading borders on rendered DOCX."""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from utils.cv_generator import cv_context_from_structured, render_cv

ROOT = Path(__file__).resolve().parent.parent
MASTER = ROOT / "templates" / "MasterTemplate.docx"

SECTIONS = ("PROFILE", "WORK EXPERIENCE", "PROJECTS", "SKILLS", "EDUCATION", "CERTIFICATIONS")


def _has_bottom_border(paragraph) -> bool:
    p_pr = paragraph._element.pPr
    return p_pr is not None and p_pr.find(qn("w:pBdr")) is not None


def _render(sample: dict) -> Document:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        out = f.name
    render_cv(str(MASTER), out, cv_context_from_structured(sample))
    return Document(out)


def test_projects_section_omitted_when_no_real_projects():
    doc = _render(
        {
            "name": "Test User",
            "contact_email": "a@b.com",
            "personal_statement": "Summary",
            "work_experience": [],
            "projects": [],
            "education": [],
            "skills": ["Python"],
            "certifications": [],
        }
    )
    assert not any(p.text.strip().upper() == "PROJECTS" for p in doc.paragraphs)


def test_work_experience_heading_has_border():
    doc = _render(
        {
            "name": "Test User",
            "contact_phone": "",
            "contact_email": "a@b.com",
            "contact_location": "",
            "linkedin": "",
            "personal_statement": "Summary",
            "work_experience": [
                {
                    "job_title": "Analyst",
                    "employer": "Acme",
                    "dates": "2023",
                    "bullets": ["Did work"],
                }
            ],
            "projects": [],
            "education": [],
            "skills": ["Python"],
            "certifications": [],
            "interests": [],
            "references": "",
        }
    )
    headings = {p.text.strip().upper(): p for p in doc.paragraphs if p.text.strip().isupper() or p.text.strip() in SECTIONS}
    for title in ("PROFILE", "WORK EXPERIENCE", "SKILLS"):
        para = next((p for p in doc.paragraphs if p.text.strip().upper() == title), None)
        assert para is not None, f"Missing {title}"
        assert _has_bottom_border(para), f"{title} missing bottom border"
